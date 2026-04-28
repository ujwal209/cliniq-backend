import os
import requests
import io
import numpy as np
import base64
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from app.core.database import get_db
import threading
import logging

logger = logging.getLogger(__name__)

_gemini_keys_raw = os.getenv("GEMINI_API_KEYS", "")
GLOBAL_GEMINI_KEYS = [k.strip() for k in _gemini_keys_raw.split(",") if k.strip()]
_gemini_lock = threading.Lock()
_gemini_index = 0

def get_rotated_gemini_keys():
    global _gemini_index
    if not GLOBAL_GEMINI_KEYS: return []
    with _gemini_lock:
        start = _gemini_index
        _gemini_index = (_gemini_index + 1) % len(GLOBAL_GEMINI_KEYS)
    return GLOBAL_GEMINI_KEYS[start:] + GLOBAL_GEMINI_KEYS[:start]

def get_embedding_model():
    keys = get_rotated_gemini_keys()
    key = keys[0] if keys else os.getenv("GEMINI_API_KEY")
    return GoogleGenerativeAIEmbeddings(model="gemini-embedding-001", google_api_key=key)

def cosine_similarity(a, b):
    return np.dot(a, b) / (np.linalg.norm(a) * np.linalg.norm(b))

async def analyze_medical_image_text(image_url: str) -> str:
    """Uses Gemini Vision to get a text description of a medical image for indexing."""
    from app.routes.vision import fetch_and_encode, get_round_robin_keys
    
    keys = get_round_robin_keys()
    if not keys: return "Failed to analyze image: No API keys."
    
    try:
        file_bytes, mime_type = fetch_and_encode(image_url)
        b64_data = base64.b64encode(file_bytes).decode('utf-8')
        
        prompt = (
            "You are a medical imaging AI. Analyze this X-ray/Medical image in extreme detail. "
            "Identify anatomical structures, abnormalities, findings, and diagnostic impressions. "
            "Provide a comprehensive textual report of what is visible so it can be indexed for future search."
        )
        
        payload = {
            "contents": [{
                "parts": [
                    {"text": prompt},
                    {"inline_data": {"mime_type": mime_type or "image/jpeg", "data": b64_data}}
                ]
            }]
        }
        
        model = "gemini-1.5-flash" # Use flash for indexing speed
        resp = requests.post(
            f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={keys[0]}",
            headers={"Content-Type": "application/json"},
            json=payload,
            timeout=30
        )
        resp.raise_for_status()
        return resp.json().get("candidates", [{}])[0].get("content", {}).get("parts", [{}])[0].get("text", "")
    except Exception as e:
        logger.error(f"Error in image analysis for indexing: {e}")
        return f"Error analyzing image: {str(e)}"

async def index_medical_report(file_url: str, patient_email: str, report_id: str, file_type: str):
    """Indices a report (PDF, Doc, or Image) into the patient's vector store."""
    db = get_db()
    text = ""
    
    try:
        if file_type == "image":
            text = await analyze_medical_image_text(file_url)
        else:
            resp = requests.get(file_url, timeout=15)
            resp.raise_for_status()
            file_bytes = resp.content
            
            if file_type == "pdf" or file_url.lower().endswith(".pdf"):
                import fitz
                doc = fitz.open(stream=file_bytes, filetype="pdf")
                for page in doc:
                    text += page.get_text() + "\n"
            elif file_type == "docx" or file_url.lower().endswith(".docx"):
                import docx
                doc = docx.Document(io.BytesIO(file_bytes))
                for para in doc.paragraphs:
                    text += para.text + "\n"
            else:
                text = file_bytes.decode('utf-8', errors='ignore')

        if not text.strip():
            logger.warning(f"No text extracted for report {report_id}")
            return

        # Split and Embed
        splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=150)
        chunks = splitter.split_text(text)
        
        embedder = get_embedding_model()
        embeddings = embedder.embed_documents(chunks)
        
        docs_to_insert = []
        for chunk, emb in zip(chunks, embeddings):
            docs_to_insert.append({
                "patient_email": patient_email,
                "report_id": report_id,
                "text": chunk,
                "embedding": emb,
                "source": file_url
            })
            
        if docs_to_insert:
            await db.medical_report_embeddings.insert_many(docs_to_insert)
            logger.info(f"Indexed {len(docs_to_insert)} chunks for report {report_id}")
            
    except Exception as e:
        logger.error(f"Failed to index report {report_id}: {e}")

async def search_patient_reports(query: str, patient_email: str, top_k: int = 5) -> str:
    """Vector search across all reports for a specific patient."""
    db = get_db()
    
    # In a production environment with millions of chunks, we'd use MongoDB's native $vectorSearch.
    # For this scale, in-memory cosine similarity is fine and easier to set up without Atlas Search Indexing.
    try:
        cursor = db.medical_report_embeddings.find({"patient_email": patient_email})
        chunks = await cursor.to_list(length=1000)
        
        if not chunks:
            return "No medical reports have been indexed yet."
            
        embedder = get_embedding_model()
        query_emb = embedder.embed_query(query)
        
        scored_chunks = []
        for c in chunks:
            score = cosine_similarity(query_emb, c["embedding"])
            scored_chunks.append((score, c["text"], c["source"]))
            
        scored_chunks.sort(key=lambda x: x[0], reverse=True)
        top_results = scored_chunks[:top_k]
        
        context = []
        for score, text, source in top_results:
            if score > 0.55: # Slightly lower threshold for medical terms
                context.append(f"[Ref: {source}]\n{text}")
                
        return "\n\n---\n\n".join(context) if context else "No relevant medical context found in your reports."
    except Exception as e:
        logger.error(f"Search error: {e}")
        return f"Error searching reports: {str(e)}"

async def index_document(doc_url: str, session_id: str):
    """Download doc, extract text, embed, and store in MongoDB for the session."""
    try:
        resp = requests.get(doc_url, timeout=15)
        resp.raise_for_status()
        file_bytes = resp.content
    except Exception as e:
        raise ValueError(f"Failed to download document: {e}")

    text = ""
    filename = doc_url.split("/")[-1].lower()

    if "pdf" in filename or doc_url.lower().endswith(".pdf"):
        import fitz
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for page in doc:
            text += page.get_text() + "\n"
    elif "docx" in filename or doc_url.lower().endswith(".docx"):
        import docx
        doc = docx.Document(io.BytesIO(file_bytes))
        for para in doc.paragraphs:
            text += para.text + "\n"
    else:
        try:
            text = file_bytes.decode('utf-8')
        except UnicodeDecodeError:
            raise ValueError("Unsupported document format or unable to decode text.")

    if not text.strip():
        raise ValueError("No extractable text found in document.")

    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    chunks = splitter.split_text(text)
    embedder = get_embedding_model()
    embeddings = embedder.embed_documents(chunks)

    db = get_db()
    docs_to_insert = []
    for chunk, emb in zip(chunks, embeddings):
        docs_to_insert.append({
            "session_id": session_id,
            "text": chunk,
            "embedding": emb,
            "source": doc_url
        })
    if docs_to_insert:
        await db.document_embeddings.insert_many(docs_to_insert)
    return len(docs_to_insert)

def search_session_documents(query: str, session_id: str, top_k: int = 3) -> str:
    """Vector search inside MongoDB using in-memory cosine similarity."""
    from pymongo import MongoClient
    client = MongoClient(os.getenv("MONGO_URI"))
    db = client[os.getenv("DB_NAME", "healthsync_db")]
    try:
        chunks = list(db.document_embeddings.find({"session_id": session_id}).limit(1000))
        if not chunks: return "No documents have been uploaded to this session yet."
        embedder = get_embedding_model()
        query_emb = embedder.embed_query(query)
        scored_chunks = []
        for c in chunks:
            score = cosine_similarity(query_emb, c["embedding"])
            scored_chunks.append((score, c["text"], c["source"]))
        scored_chunks.sort(key=lambda x: x[0], reverse=True)
        top_results = scored_chunks[:top_k]
        context = []
        for score, text, source in top_results:
            if score > 0.6:
                context.append(f"[Source: {source}]\n...{text}...")
        return "\n\n---\n\n".join(context) if context else "No relevant info found."
    finally:
        client.close()
